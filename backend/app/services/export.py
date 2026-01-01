"""Export service for clinical notes."""

import io
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


def export_to_pdf(content: str, title: str = "Clinical Note") -> bytes:
    """
    Export clinical note content to PDF.
    
    Args:
        content: The note content (plain text or simple markdown)
        title: Document title
        
    Returns:
        PDF file as bytes
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
    )

    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Heading1"],
        fontSize=16,
        spaceAfter=20,
        textColor=colors.darkblue,
    )
    
    body_style = ParagraphStyle(
        "CustomBody",
        parent=styles["Normal"],
        fontSize=11,
        leading=14,
        spaceAfter=8,
    )
    
    header_style = ParagraphStyle(
        "CustomHeader",
        parent=styles["Heading2"],
        fontSize=13,
        spaceBefore=12,
        spaceAfter=6,
        textColor=colors.darkblue,
    )

    story = []

    # Title
    story.append(Paragraph(title, title_style))
    story.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", styles["Normal"]))
    story.append(Spacer(1, 20))

    # Process content - handle simple sections
    lines = content.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            story.append(Spacer(1, 8))
        elif line.startswith("##"):
            # Section header
            story.append(Paragraph(line.replace("#", "").strip(), header_style))
        elif line.startswith("#"):
            # Main header
            story.append(Paragraph(line.replace("#", "").strip(), title_style))
        elif line.startswith("•") or line.startswith("-"):
            # Bullet point
            story.append(Paragraph(f"&bull; {line[1:].strip()}", body_style))
        else:
            story.append(Paragraph(line, body_style))

    doc.build(story)
    return buffer.getvalue()


def export_to_docx(content: str, title: str = "Clinical Note") -> bytes:
    """
    Export clinical note content to DOCX.
    
    Args:
        content: The note content (plain text or simple markdown)
        title: Document title
        
    Returns:
        DOCX file as bytes
    """
    document = Document()

    # Set default font
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    document.add_heading(title, 0)
    document.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    document.add_paragraph()

    # Process content
    lines = content.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            document.add_paragraph()
        elif line.startswith("##"):
            # Section header
            document.add_heading(line.replace("#", "").strip(), level=2)
        elif line.startswith("#"):
            # Main header
            document.add_heading(line.replace("#", "").strip(), level=1)
        elif line.startswith("•") or line.startswith("-"):
            # Bullet point
            p = document.add_paragraph(style="List Bullet")
            p.add_run(line[1:].strip())
        elif line.startswith(tuple("0123456789")):
            # Numbered item
            p = document.add_paragraph(style="List Number")
            # Remove leading number and punctuation
            text = line.lstrip("0123456789").lstrip(".").strip()
            p.add_run(text)
        else:
            document.add_paragraph(line)

    # Save to buffer
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()

